# encoding: utf-8
from __future__ import absolute_import, unicode_literals
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text import Paragraph, Run


__author__ = 'bluec0re'


def _insert_after(self, text=None, style=None):
    p = OxmlElement('w:p')
    self._p.addnext(p)
    paragraph = Paragraph(p, self._parent)
    if text:
        paragraph.add_run(text)
    if style is not None:
        paragraph.style = style
    return paragraph

Paragraph.insert_paragraph_after = _insert_after


def _add_caption(self, text, sequence_name=None):
    if not sequence_name:
        sequence_name = 'Figure'

    paragraph = self.insert_paragraph_after(text=None, style='Caption')
    paragraph.add_run('%s ' % sequence_name)
    new_fld = OxmlElement('w:fldSimple', attrs={
        qn('w:instr'): r' SEQ %s \* ARABIC ' % sequence_name
    })
    new_r = OxmlElement('w:r')
    new_r.add_t('0')
    new_fld.append(new_r)
    paragraph._p.append(new_fld)
    paragraph.add_run(": %s" % text)

    return paragraph


Paragraph.add_caption = _add_caption


def _append_run(self, run, text, style=None):
    new_r = run._r.add_r_after()
    new_r = Run(new_r, self)
    if text:
        new_r.text = text
    if style:
        new_r.style = style
    return new_r

Paragraph.append_run = _append_run
