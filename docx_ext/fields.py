# encoding: utf-8
from __future__ import absolute_import, unicode_literals

from collections import defaultdict
from functools import partial
from io import BytesIO
import shlex
import logging
import re

from docx.oxml.ns import qn
from docx.shared import Parented
from docx.table import Table
from docx.text import Paragraph, Run
from lxml.etree import QName
import lxml.html
from docx_ext.parser import ParserException
from docx_ext.textbox import Textbox


try:
    from PIL.Image import Image
except ImportError:
    Image = None

from .utils import make_abs, make_relative


ALLOWED_TAGS = (
    'p',
    'span',
    'br',
    'img',
    'i',
    'b',
    'h1',
    'h2',
    'h3',
    'h4'
)

ALLOWED_ATTRS = {
    '*': (
        'class',
        'style'
    ),
    'img': (
        'alt',
        'src'
    )
}

ALLOWED_STYLES = (
    'color',
    'font-weight',
    'font-style',
    'font-variant'
)

__author__ = 'bluec0re'

log = logging.getLogger(__name__)


def clean_style(styles):
    """
    >>> clean_style("color: #ff0000")
    u'color: #ff0000'
    >>> clean_style("color:  #ff0000 ")
    u'color: #ff0000'
    >>> clean_style("color:  #ff0000; foo: bar")
    u'color: #ff0000'
    >>> clean_style("bar: foo; color:#ff0000; foo: bar")
    u'color: #ff0000'
    """
    result = []
    styles = styles.split(';')
    for style in styles:
        name, value = style.split(':', 1)
        name = name.strip()
        value = value.strip()
        if name in ALLOWED_STYLES:
            result.append('{0}: {1}'.format(name, value))
    return '; '.join(result)


def clean_attr(allowed_attrs, m):
    name = m.group('name').strip().lower()
    if name in allowed_attrs:
        if name == 'style':
            return ' {0}="{1}"'.format(name, clean_style(m.group('value')))
        else:
            return m.group(0)
    return ''


def clean_attrs(m):
    # noinspection PyTypeChecker
    return re.sub(r'\s+(?P<name>[^ =]+)\s*=\s*"?(?P<value>[^"]+)"?',
                  partial(clean_attr, ALLOWED_ATTRS.get(m.group('name'), ()) + ALLOWED_ATTRS['*']),
                  m.group(0))


def clean_html(html):
    """
    >>> clean_html('test')
    u'test'
    >>> clean_html('<span>test</span>')
    u'<span>test</span>'
    >>> clean_html('test <iframe>a</iframe>')
    u'test a'
    >>> clean_html('<span class="Heading1">test</span>')
    u'<span class="Heading1">test</span>'
    >>> clean_html('<span class="Heading1" id="foo">test</span>')
    u'<span class="Heading1">test</span>'
    >>> clean_html('<span style="color: #ff0000">test</span>')
    u'<span style="color: #ff0000">test</span>'
    >>> clean_html('<span style="color: #ff0000; background: foo">test</span>')
    u'<span style="color: #ff0000">test</span>'
    >>> clean_html('<script style="color: #ff0000; background: foo">test</foo>')
    u'test'
    >>> clean_html('<a>test</a>')
    u'test'
    >>> clean_html('<i>test</i>')
    u'<i>test</i>'
    >>> clean_html('<b>test</b>')
    u'<b>test</b>'
    >>> clean_html('<img/>test<img />')
    u'<img/>test<img />'
    """
    tags = ['({0}[ >/])'.format(re.escape(tag)) for tag in ALLOWED_TAGS]
    tags += ['(/{0}[ >/])'.format(re.escape(tag)) for tag in ALLOWED_TAGS]
    regex = '<(?!{0})[^>]*>'.format('|'.join(tags))
    # print(regex)
    html = re.sub(regex, '', html)

    html = re.sub(r'<(?P<name>[^ >]+)[^>]*>', clean_attrs, html)

    return html


_IMAGE_FACTORIES = set([])


def image_factory(path):
    for fact in _IMAGE_FACTORIES:
        img = fact(path)
        if img is not None:
            return img


def register_image_factory(factory):
    global _IMAGE_FACTORIES
    _IMAGE_FACTORIES.add(factory)


def unregister_image_factory(factory):
    global _IMAGE_FACTORIES
    _IMAGE_FACTORIES.erase(factory)


# noinspection PyProtectedMember
class Field(Parented):
    def __init__(self, parent, code=None, default=None):
        super(Field, self).__init__(parent)
        self.code = code
        self.default = default
        self.format = defaultdict(list)
        self.extra = []
        self.__xpath_start = None
        self.__xpath_end = None
        self.paragraph = None
        self.__start = None
        self.__end = None
        self._base = None

    def __unicode__(self):
        if self.default:
            return "%s (%s)" % (self.code, self.default)
        else:
            return self.code

    def __str__(self):
        return unicode(self).encode('utf-8')

    def __repr__(self):
        return ("%s(code=%s, default=%s, format=%s, extra=%s, from=%s, to=%s)" % (
            type(self).__name__,
            self.code,
            self.default,
            self.format,
            self.extra,
            self.xpath_start,
            self.xpath_end
        ))#.encode('utf-8')

    @property
    def xpath_start(self):
        if self.__start is not None:
            return self.__start.getroottree().getpath(self.__start)
        else:
            return self.__xpath_start

    @xpath_start.setter
    def xpath_start(self, value):
        self.__xpath_start = value
        self.__start = None
        assert self.start is not None

    @property
    def xpath_end(self):
        if self.__end is not None:
            return self.__end.getroottree().getpath(self.__end)
        else:
            return self.__xpath_end

    @xpath_end.setter
    def xpath_end(self, value):
        self.__xpath_end = value
        self.__end = None
        assert self.end is not None or self.__start is not None

    @property
    def start(self):
        base = self._base if self._base is not None else self._parent._p
        if self.__start is None and self.xpath_start:
            log.debug("Getting start %s", self.xpath_start)
            try:
                self.__start = base.xpath(self.xpath_start)[0]
            except IndexError:
                log.warning("Couldn't find start: %s", self.xpath_start)

        if self.__start is not None and self.__end is not None and \
           self.__start.getparent() == self.__end.getparent():
            self._parent._p = self.__start.getparent()
        return self.__start

    @start.setter
    def start(self, value):
        self.__start = value
        if value is not None:
            self.__xpath_start = self.__start.getroottree().getpath(value)

    @property
    def end(self):
        base = self._base if self._base is not None else self._parent._p
        if self.__end is None and self.xpath_end:
            log.debug("Getting end %s", self.xpath_end)
            try:
                self.__end = base.xpath(self.xpath_end)[0]
            except IndexError:
                log.warning("Couldn't find end: %s", self.xpath_end)

        if self.__start is not None and self.__end is not None and \
           self.__start.getparent() == self.__end.getparent():
            self._parent._p = self.__start.getparent()
        return self.__end

    @end.setter
    def end(self, value):
        self.__end = value
        if value is not None:
            self.__xpath_end = self.__end.getroottree().getpath(value)

    def insert(self, run, obj, allowed_styles=None):
        if obj is None:
            return

        if Image is not None and isinstance(obj, Image):
            io = BytesIO()
            obj.save(io, 'PNG')
            doc = run._parent._parent._parent
            # TODO: verify correctness
            section = doc.sections[-1]
            width = section.page_width - section.left_margin - section.right_margin
            run.add_picture(io, width=width)

            if hasattr(obj, 'caption'):
                para = run._parent
                return para.add_caption(obj.caption)
        elif re.search('<(.+?)>', obj):
            html = clean_html(obj)
            root = lxml.html.fromstring(html)

            def _get_styles(elem):
                styles = {}
                if elem.tag == 'i':
                    styles['font-style'] = 'italic'
                elif elem.tag == 'b':
                    styles['font-weight'] = 'bold'

                if 'style' not in elem.attrib:
                    return styles

                _styles = elem.attrib['style'].split(';')

                styles.update({
                    style.split(':', 1)[0].strip(): style.split(':', 1)[1].strip() for style in _styles
                })

                return styles

            def _get_class(elem):
                clazz = elem.attrib.get('class')
                if not clazz:
                    if elem.tag == 'h1':
                        clazz = 'Heading1'
                    elif elem.tag == 'h2':
                        clazz = 'Heading2'
                    elif elem.tag == 'h3':
                        clazz = 'Heading3'
                    elif elem.tag == 'h4':
                        clazz = 'Heading4'

                if clazz and allowed_styles and clazz not in allowed_styles:
                    msg = "Style %s does not exist in given template" % clazz
                    log.critical(msg)
                    raise ValueError(msg)
                return clazz

            def _transform(currentrun, el, is_first=False):
                log.debug('Transforming element %s into %s', el, currentrun._r.getroottree().getpath(currentrun._r))
                parts = [el.text] + [x for c in el for x in (c, c.tail)]
                if isinstance(parts[0], str):
                    currentrun.text = parts[0]

                for part in parts[1:]:
                    if part is None:
                        continue
                    log.debug("Run path: %s", currentrun._r.getroottree().getpath(currentrun._r))
                    if isinstance(part, str):
                        currentrun = currentrun._parent.append_run(currentrun, part)
                    elif part.tag == 'p':
                        if not is_first:
                            log.debug('New paragraph')
                            p = currentrun._parent.insert_paragraph_after()
                            p.style = _get_class(part)
                            currentrun = p.add_run()
                        is_first = False
                        currentrun = _transform(currentrun, part)
                    else:
                        clazz = _get_class(part)
                        if clazz and clazz.startswith('Heading'):
                            log.debug('New heading')
                            # create new paragraph
                            p = currentrun._parent.insert_paragraph_after()
                            p.style = clazz
                            currentrun = p.add_run()
                        else:
                            currentrun = currentrun._parent.append_run(currentrun, None, style=clazz)
                        styles = _get_styles(part)
                        if 'color' in styles:
                            currentrun.color = styles['color'].replace('#', '')
                        if 'font-weight' in styles:
                            currentrun.bold = 'bold' in styles['font-weight']
                        if 'font-variant' in styles:
                            currentrun.small_caps = 'small-caps' in styles['font-variant']
                        if 'font-style' in styles:
                            currentrun.italic = 'italic' in styles['font-style']

                        if part.tag == 'img':
                            log.debug('New image')
                            img = image_factory(part.attrib.get('src'))
                            if img is not None:
                                img.caption = part.attrib.get('alt')
                            p = currentrun._parent.insert_paragraph_after()
                            currentrun = p.add_run()
                            p = self.insert(currentrun, img)
                            currentrun = p.runs[-1]
                            # currentrun = p.add_run()  # maybe not optimal
                        else:
                            currentrun = _transform(currentrun, part)
                return currentrun

            _transform(run, root, True)
        else:
            run.text = obj

    def replace(self, text, base=None, allowed_styles=None):
        log.debug("Replacing content from %s with '%s' in %s", self, text, base)
        if base is None or True:
            start, end = self.start, self.end
        else:
            start = base.xpath(self.xpath_start)[0]
            end = base.xpath(self.xpath_end)[0]

        log.debug("Using start: %s %s %s", start, self.xpath_start, start.getparent())
        log.debug("Using end: %s %s %s", end, self.xpath_end, end.getparent() if end is not None else None)

        for sibl in start.itersiblings():
            sibl.getparent().remove(sibl)
            if [sibl] == end:
                break

        if start.tag.endswith('r'):
            r = Run(start, self._parent)
            if hasattr(r._r, 'clear_content'):
                r._r.clear_content()
            self.insert(r, text, allowed_styles=allowed_styles)
            return r
        else:
            # TODO
            pass

    def remove(self):
        r = self.replace('')
        if r is None:
            return

        p = r._parent
        parent = r._r.getparent()
        log.debug("%s P: %s", self, parent)
        if parent is not None:
            parent.remove(r._r)
        else:
            log.warning("Parent is none for %s", self)

        if p._p is None:
            log.warning("Parent of %s has no element attached to it", self)
            return

        if len(p.runs) == 0:
            parent = p._p.getparent()
            if parent is not None:
                parent.remove(p._p)
            else:
                log.warning("No parent for paragraph %s", p)
                log.debug("Text: %s, Style: %s", p.text, p.style)

    def update_xpaths(self, root_xpath, base_xpath):
        self.__start = None
        xp = make_relative(self.xpath_start, root_xpath)
        xp = make_abs(xp, base_xpath)
        self.xpath_start = xp

        self.__end = None
        xp = make_relative(self.xpath_end, root_xpath)
        xp = make_abs(xp, base_xpath)
        self.xpath_end = xp


# noinspection PyProtectedMember
def _fields(self):
    fields = []
    field = None
    # noinspection PyPep8Naming
    fldCharType = qn('w:fldCharType')
    instr = qn('w:instr')
    root = self._p.getroottree()
    instructions = ''

    for run in self.runs:
        for chld in run._r.getchildren():
            tag = QName(chld.tag).localname
            if tag == 'fldChar':
                if chld.attrib.get(fldCharType) == 'begin':
                    field = Field(self)
                    field.xpath_start = root.getpath(run._r)
                elif chld.attrib.get(fldCharType) == 'end' and field is not None:
                    if instructions:
                        lex = shlex.shlex(instructions, posix=True)
                        lex.whitespace_split = True
                        lex.commenters = ''
                        lex.escape = ''
                        try:
                            tokens = list(lex)
                        except ValueError as e:
                            raise ParserException(str(e) + chld.text, chld)

                        field.code = tokens[0]
                        fmt_target = None
                        for token in tokens[1:]:
                            if fmt_target is not None:
                                field.format[fmt_target].append(token)
                                fmt_target = None
                            elif token.startswith('\\'):
                                fmt_target = token[1:]
                            else:
                                field.extra.append(token)
                    field.xpath_end = root.getpath(run._r)
                    fields.append(field)
                    field = None
                    instructions = ''
            elif field:
                if tag == 't':
                    field.default = chld.text
                elif tag == 'instrText' and chld.text.strip():
                    instructions += chld.text

    for fld in self._p.xpath('./w:fldSimple'):
        if fld.attrib.get(instr):
            field = Field(self)
            field.xpath_start = root.getpath(fld)
            field.xpath_end = root.getpath(fld)
            if fld.find(qn('w:r')):
                field.default = fld.find(qn('w:r'))[0].text
            lex = shlex.shlex(fld.attrib.get(instr), posix=True)
            lex.whitespace_split = True
            lex.commenters = ''
            lex.escape = ''
            tokens = list(lex)
            field.code = tokens[0]
            fmt_target = None
            for token in tokens[1:]:
                if fmt_target is not None:
                    field.format[fmt_target].append(token)
                    fmt_target = None
                elif token.startswith('\\'):
                    fmt_target = token[1:]
                else:
                    field.extra.append(token)
            fields.append(field)
    return fields


Paragraph.fields = _fields


# noinspection PyProtectedMember
def _fields(self):
    fields = []

    for row in self.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                fields += paragraph.fields()
    return fields

Table.fields = _fields


# noinspection PyProtectedMember
def _fields(self):
    fields = []

    for paragraph in self.paragraphs:
        fields += paragraph.fields()

    for table in self.tables:
        fields += table.fields()

    return fields

Textbox.fields = _fields
